import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import pandas as pd
from io import BytesIO
import datetime
import math
import sys

# Intentamos importar python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ==============================================================================
# 1. BASE DE PRECIOS UNITARIOS (ORIGINAL)
# ==============================================================================
PRECIOS_REF = {
    # --- OBRA CIVIL / GENERAL ---
    "Adecuación Sala/Obra Civil (m2)": 850.0,
    "Refuerzo Estructural (m2)": 150.0,
    "Suelo Técnico (m2)": 120.0,
    "Cerramiento/Contención (ud)": 3500.0,
    
    # --- ELÉCTRICO ---
    "Celda MT (ud)": 18000.0,
    "Trafo 1000-2500kVA (ud)": 45000.0,
    "Generador Diesel (kVA)": 200.0, 
    "UPS Modular (kW)": 250.0, 
    "CGBT (ud)": 25000.0,
    "Cuadro Distribución IT (ud)": 8000.0,
    "Blindobarra (m)": 450.0,
    "Cableado Potencia Grueso (m)": 60.0, 
    "Cableado Potencia Medio (m)": 25.0,  
    "Cableado Rack (ud)": 50.0, 
    "Bandeja Eléctrica (m)": 45.0,
    
    # --- CLIMA ---
    "Chiller (kW)": 150.0, 
    "CRAH/InRow (ud)": 18000.0,
    "Tubería Acero DN100-200 (m)": 180.0, 
    "Tubería Cobre/PPR Pequeña (m)": 45.0,
    "Válvulas y Accesorios (% Tubería)": 0.30, 
    "Bomba Circuladora (ud)": 4500.0,
    
    # --- DLC ---
    "CDU (ud)": 35000.0,
    "Manifold Rack (ud)": 2000.0,
    "Latiguillos DLC (ud)": 150.0,
    
    # --- PCI & SEGURIDAD ---
    "Centralita Incendios (ud)": 2500.0,
    "Detector/Sensor (ud)": 150.0,
    "Cilindro NOVEC 1230 (Kg)": 60.0, 
    "Cilindro ARGONITE (m3)": 40.0,
    "Grupo Bombeo Nebulizada (ud)": 40000.0, 
    "Boquilla Nebulizada (ud)": 200.0,
    "Cámara CCTV (ud)": 400.0,
    "Control Acceso (punto)": 1200.0,
    
    # --- COMUNICACIONES & BMS ---
    "Rack 42U (ud)": 1200.0,
    "Fibra Óptica OM4/OS2 (m)": 8.0,
    "Cable Cobre Cat6A (m)": 3.0, 
    "Bandeja Rejilla/Fibra (m)": 40.0,
    "Punto BMS/Integración (ud)": 350.0
}

# ==============================================================================
# 2. CLASE PRINCIPAL: MOTOR DE CÁLCULO (TU CÓDIGO EXACTO)
# ==============================================================================

class DisenadorV14:
    def __init__(self, redundancia_electrica, redundancia_hvac, suministro_AB, distribucion_IT_tipo,  
                 # Carga y Diseño
                 num_cerramientos, racks_por_cerramiento, servidores_por_rack, tipo_cerramiento, 
                 P_idle, P_max, 
                 # Parámetros HVAC/DLC
                 P_iluminacion, P_otras_fuerza, 
                 cop_hvac_aire, T_entrada_aire, T_salida_aire, 
                 prodfrio_tec, intcalor_tec, distribfrio_tec, n_intercambiadores, 
                 # Parámetros DLC Detallados 
                 cerramientos_con_dlc, tipo_gen_frio_dlc, cop_dlc_gen, 
                 tipo_dist_frio_dlc, pot_aux_dlc_dist, eficiencia_captura_dlc, 
                 # Parámetros Auxiliares
                 centralitas_incendios, vesda_unidades, grupos_bombeo_pci, cctv_unidades, control_accesos_pax,
                 # Parámetros V13
                 tecnologia_pci,
                 # NUEVOS PARÁMETROS V14 (DIMENSIONALES)
                 num_plantas, area_por_planta, area_sala_it):
        
        # --- Datos Dimensionales ---
        self.num_plantas = num_plantas
        self.area_por_planta = area_por_planta
        self.area_sala_it = area_sala_it
        self.area_total_construida = num_plantas * area_por_planta
        self.altura_planta = 4.5 
        
        # --- Datos de Entrada Previos ---
        self.tecnologia_pci = tecnologia_pci 
        self.centralitas_incendios = centralitas_incendios
        self.vesda_unidades = vesda_unidades
        self.grupos_bombeo_pci = grupos_bombeo_pci
        self.cctv_unidades = cctv_unidades
        self.control_accesos_pax = control_accesos_pax
        self.P_iluminacion = P_iluminacion
        self.P_otras_fuerza = P_otras_fuerza
        
        # Maquinaria
        self.prodfrio_tec = prodfrio_tec
        self.intcalor_tec = intcalor_tec
        self.distribfrio_tec = distribfrio_tec
        self.n_intercambiadores = n_intercambiadores
        self.T_entrada_aire = T_entrada_aire
        self.T_salida_aire = T_salida_aire

        # Carga
        self.servidores_por_rack = servidores_por_rack
        self.P_max_servidor = P_max
        self.N_servidores_total = num_cerramientos * racks_por_cerramiento * servidores_por_rack
        self.P_IT_demandada = self.N_servidores_total * P_max 
        self.num_cerramientos = num_cerramientos
        self.racks_por_cerramiento = racks_por_cerramiento
        self.num_racks_total = num_cerramientos * racks_por_cerramiento
        self.tipo_cerramiento = tipo_cerramiento
        self.P_IT_por_rack = servidores_por_rack * P_max 

        # Redundancia
        self.R_elec = redundancia_electrica
        self.R_hvac = redundancia_hvac
        self.factor_N_elec = self._get_factor_redundancia(self.R_elec)
        self.factor_N_hvac = self._get_factor_redundancia(self.R_hvac)
        
        # Eléctrico / HVAC / DLC
        self.Suministro_AB = suministro_AB
        self.Distribucion_IT_tipo = distribucion_IT_tipo
        self.COP_HVAC = cop_hvac_aire
        self.COP_DLC_GEN = cop_dlc_gen
        self.P_DLC_dist_por_cerr = pot_aux_dlc_dist
        self.Eficiencia_Captura_DLC = eficiencia_captura_dlc
        self.cerramientos_con_dlc = cerramientos_con_dlc
        self.tipo_gen_frio_dlc = tipo_gen_frio_dlc
        self.tipo_dist_frio_dlc = tipo_dist_frio_dlc
        
        # Cálculos Potencia
        self.P_PCI_calc = (grupos_bombeo_pci * 20000) + (centralitas_incendios * 500)
        self.P_Control_calc = (cctv_unidades * 100) + (control_accesos_pax * 50) + (vesda_unidades * 150)
        self.P_HVAC_demandada, self.P_DLC_demandada = self._calcular_cargas_electricas_refrigeracion()
        self.P_Aux_total = self.P_iluminacion + self.P_otras_fuerza + self.P_PCI_calc + self.P_Control_calc
        self.P_total_demandada = self.P_IT_demandada + self.P_HVAC_demandada + self.P_DLC_demandada + self.P_Aux_total

    def _get_factor_redundancia(self, r):
        if r == "N": return 1.0
        if r == "N+1": return 1.25 
        if r == "2N": return 2.0
        if r == "2N+1": return 2.25
        return 1.0
    
    def _calcular_cargas_electricas_refrigeracion(self):
        Q_DLC_capturada = self.P_IT_demandada * (self.cerramientos_con_dlc / self.num_cerramientos) * self.Eficiencia_Captura_DLC
        P_DLC_gen = Q_DLC_capturada / self.COP_DLC_GEN if self.COP_DLC_GEN > 0 else 0
        P_DLC_dist = self.cerramientos_con_dlc * self.P_DLC_dist_por_cerr
        P_DLC_demandada = P_DLC_gen + P_DLC_dist

        Q_Remanente = self.P_IT_demandada - Q_DLC_capturada
        factor_eficiencia_aire = 1.05 if self.tipo_cerramiento == "Pasillo Frío" else 1.25
        Q_HVAC_aire_requerida = Q_Remanente * factor_eficiencia_aire
        P_HVAC_demandada = Q_HVAC_aire_requerida / self.COP_HVAC if self.COP_HVAC > 0 else 0
        
        return P_HVAC_demandada, P_DLC_demandada

    # --- MOTOR HIDRÁULICO ---
    def _calcular_tuberia_colector(self, Q_kW, delta_T):
        if Q_kW <= 0.1:
            return {"Caudal_Total_m3h": 0, "DN_mm": 0, "Velocidad_ms": 0, "Material": "-", "Num_Circuitos": 0, "Longitud_Estimada_m": 0}

        rho = 1000; Cp = 4.18 
        m_kgs = Q_kW / (Cp * delta_T)
        V_m3s = m_kgs / rho
        V_m3h = V_m3s * 3600
        
        diametros = [(50,"PPR/Cobre"),(65,"Acero Carb."),(80,"Acero Carb."),(100,"Acero Carb."),(125,"Acero Carb."),(150,"Acero Carb."),(200,"Acero Carb."),(250,"Acero Carb."),(300,"Acero Carb.")]
        num_circuitos = 1; seleccion = None
        
        dist_horizontal = math.sqrt(self.area_por_planta) * 1.5 
        dist_vertical = self.altura_planta * self.num_plantas
        longitud_total = (dist_horizontal + dist_vertical) * 2 
        
        while seleccion is None:
            caudal_por_circuito = V_m3s / num_circuitos
            mejor_dn = None; mejor_vel = 0.0; mejor_mat = ""
            for dn, mat in diametros:
                area = np.pi * ((dn/1000.0)**2) / 4.0
                vel = caudal_por_circuito / area
                if vel <= 2.5: 
                    mejor_dn = dn; mejor_vel = vel; mejor_mat = mat
                    break
            
            if mejor_dn is not None:
                seleccion = {
                    "Caudal_Total_m3h": V_m3h, "DN_mm": mejor_dn, "Velocidad_ms": mejor_vel, 
                    "Material": mejor_mat, "Num_Circuitos": num_circuitos,
                    "Longitud_Estimada_m": longitud_total * num_circuitos
                }
            else:
                num_circuitos += 1
                if num_circuitos > 50: seleccion = {"Caudal_Total_m3h": 0, "DN_mm": 0, "Velocidad_ms": 0, "Material": "-", "Num_Circuitos": 0, "Longitud_Estimada_m": 0}
        return seleccion

    # --- RATIOS ---
    def calcular_kpis_densidad(self, Q_inst_hvac, S_inst_elec_kVA):
        if self.area_sala_it <= 0 or self.area_total_construida <= 0: return {}
        
        densidad_it = (self.P_IT_demandada / 1000) / self.area_sala_it
        densidad_elec_instalada = S_inst_elec_kVA / self.area_total_construida
        densidad_termica_it = Q_inst_hvac / self.area_sala_it
        racks_m2 = self.num_racks_total / self.area_sala_it
        
        return {
            "Densidad Potencia IT (kW/m² IT)": densidad_it,
            "Densidad Potencia Elec. Instalada (kVA/m² Const.)": densidad_elec_instalada,
            "Densidad Térmica Refrigeración (kWth/m² IT)": densidad_termica_it,
            "Densidad Física (Racks/m² IT)": racks_m2
        }

    # --- CAPEX ESTIMATION ---
    def calcular_presupuesto_detallado(self, res_elec, res_hvac, res_dlc):
        items = []
        
        lado_planta = math.sqrt(self.area_por_planta)
        altura_total = self.num_plantas * self.altura_planta
        
        # 1. OBRA CIVIL
        items.append({"Cat": "Civil", "Item": "Adecuación Arquitectónica (Suelo/Pintura)", "Ud": "m2", "Cant": self.area_total_construida, "PU": PRECIOS_REF["Adecuación Sala/Obra Civil (m2)"]})
        items.append({"Cat": "Civil", "Item": "Suelo Técnico Elevado", "Ud": "m2", "Cant": self.area_sala_it, "PU": PRECIOS_REF["Suelo Técnico (m2)"]})
        items.append({"Cat": "Civil", "Item": "Contención Pasillos/Cerramientos", "Ud": "ud", "Cant": self.num_cerramientos, "PU": PRECIOS_REF["Cerramiento/Contención (ud)"]})
        items.append({"Cat": "Civil", "Item": "Racks Servidores", "Ud": "ud", "Cant": self.num_racks_total, "PU": PRECIOS_REF["Rack 42U (ud)"]})

        # 2. ELÉCTRICO
        lados = res_elec['Num_Lados']
        items.append({"Cat": "Eléctrico", "Item": "Celdas Media Tensión", "Ud": "ud", "Cant": res_elec['Num_Celdas_MT'], "PU": PRECIOS_REF["Celda MT (ud)"]})
        items.append({"Cat": "Eléctrico", "Item": "Transformadores", "Ud": "ud", "Cant": lados, "PU": PRECIOS_REF["Trafo 1000-2500kVA (ud)"]})
        pot_gen = res_elec['S_Total_N_kVA'] * self.factor_N_elec 
        items.append({"Cat": "Eléctrico", "Item": "Grupos Electrógenos", "Ud": "kVA", "Cant": pot_gen, "PU": PRECIOS_REF["Generador Diesel (kVA)"]})
        items.append({"Cat": "Eléctrico", "Item": "SAI / UPS", "Ud": "kW", "Cant": self.P_total_demandada/1000 * self.factor_N_elec, "PU": PRECIOS_REF["UPS Modular (kW)"]})
        items.append({"Cat": "Eléctrico", "Item": "Cuadros CGBT", "Ud": "ud", "Cant": lados, "PU": PRECIOS_REF["CGBT (ud)"]})
        
        dist_mt = (altura_total + 50) * lados 
        dist_bt_principal = 20 * lados 
        dist_promedio_sala = (altura_total / 2) + (lado_planta / 2) 
        dist_lineas_sala = dist_promedio_sala * self.num_cerramientos * lados
        
        items.append({"Cat": "Eléctrico", "Item": "Cableado MT/BT Acometida", "Ud": "m", "Cant": dist_mt + dist_bt_principal, "PU": PRECIOS_REF["Cableado Potencia Grueso (m)"]}) 
        items.append({"Cat": "Eléctrico", "Item": "Blindobarras / Líneas Sala", "Ud": "m", "Cant": dist_lineas_sala + (self.num_racks_total * 2), "PU": PRECIOS_REF["Blindobarra (m)"]})
        items.append({"Cat": "Eléctrico", "Item": "Bandejas Portacables Elec.", "Ud": "m", "Cant": dist_lineas_sala, "PU": PRECIOS_REF["Bandeja Eléctrica (m)"]})
        items.append({"Cat": "Eléctrico", "Item": "Cableado Última Milla (Rack)", "Ud": "ud", "Cant": self.num_racks_total * 2, "PU": PRECIOS_REF["Cableado Rack (ud)"]})

        # 3. CLIMATIZACIÓN (HVAC)
        q_hvac = res_hvac['Q_Instalada_kW']
        items.append({"Cat": "HVAC", "Item": "Equipos Producción (Chillers/Torres)", "Ud": "kW_frío", "Cant": q_hvac, "PU": PRECIOS_REF["Chiller (kW)"]})
        
        n_equipos_hvac = np.ceil(q_hvac / 100) 
        items.append({"Cat": "HVAC", "Item": "Equipos Sala (CRAH/InRow)", "Ud": "ud", "Cant": n_equipos_hvac, "PU": PRECIOS_REF["CRAH/InRow (ud)"]})
        
        len_hvac = res_hvac["Hidro_Prim"]["Longitud_Estimada_m"] + res_hvac["Hidro_Sec"]["Longitud_Estimada_m"]
        coste_tubo_hvac = len_hvac * PRECIOS_REF["Tubería Acero DN100-200 (m)"]
        items.append({"Cat": "HVAC", "Item": "Tuberías Acero (Aisladas)", "Ud": "m", "Cant": len_hvac, "PU": PRECIOS_REF["Tubería Acero DN100-200 (m)"]})
        items.append({"Cat": "HVAC", "Item": "Válvulas, Bombas y Accesorios", "Ud": "Global", "Cant": 1, "PU": coste_tubo_hvac * 0.4 + (PRECIOS_REF["Bomba Circuladora (ud)"]*4)})

        # 4. DLC 
        if self.cerramientos_con_dlc > 0:
            q_dlc = res_dlc['Q_DLC_kW']
            items.append({"Cat": "DLC", "Item": "CDUs (Coolant Distribution Units)", "Ud": "ud", "Cant": self.cerramientos_con_dlc, "PU": PRECIOS_REF["CDU (ud)"]})
            len_dlc = res_dlc["Hidro_Prim"]["Longitud_Estimada_m"] + res_dlc["Hidro_Sec"]["Longitud_Estimada_m"]
            items.append({"Cat": "DLC", "Item": "Red Hidráulica DLC", "Ud": "m", "Cant": len_dlc, "PU": PRECIOS_REF["Tubería Cobre/PPR Pequeña (m)"]})
            items.append({"Cat": "DLC", "Item": "Manifolds & Latiguillos Rack", "Ud": "ud", "Cant": self.cerramientos_con_dlc * self.racks_por_cerramiento, "PU": PRECIOS_REF["Manifold Rack (ud)"]})

        # 5. PCI 
        volumen_total_construido = self.area_total_construida * self.altura_planta
        volumen_sala_it = self.area_sala_it * self.altura_planta
        
        items.append({"Cat": "PCI", "Item": "Sistema Detección (Central+Sensores)", "Ud": "ud", "Cant": 1, "PU": PRECIOS_REF["Centralita Incendios (ud)"] + (self.num_racks_total * PRECIOS_REF["Detector/Sensor (ud)"])})
        
        if self.tecnologia_pci == "Agua Nebulizada":
            items.append({"Cat": "PCI", "Item": "Grupo Bombeo Nebulizada", "Ud": "ud", "Cant": 1, "PU": PRECIOS_REF["Grupo Bombeo Nebulizada (ud)"]})
            metros_tubo_pci = math.sqrt(self.area_total_construida) * self.num_plantas * 2 
            items.append({"Cat": "PCI", "Item": "Red Tubería Inox + Boquillas", "Ud": "ud", "Cant": int(self.area_total_construida/20), "PU": PRECIOS_REF["Boquilla Nebulizada (ud)"] * 3}) 
        elif self.tecnologia_pci == "NOVEC 1230":
            kg_novec = volumen_sala_it * 0.75 
            items.append({"Cat": "PCI", "Item": "Gas NOVEC 1230 (Sala IT)", "Ud": "Kg", "Cant": kg_novec, "PU": PRECIOS_REF["Cilindro NOVEC 1230 (Kg)"]})
        else: 
            m3_gas = volumen_sala_it * 0.5 
            items.append({"Cat": "PCI", "Item": "Cilindros Gas Inerte (Sala IT)", "Ud": "m3", "Cant": m3_gas, "PU": PRECIOS_REF["Cilindro ARGONITE (m3)"]})

        # 6. COMUNICACIONES 
        n_servers = self.N_servidores_total
        backbone_fibra = altura_total * 4 
        horizontal_fibra = (math.sqrt(self.area_sala_it) + 10) * self.num_racks_total 
        total_fibra = backbone_fibra + horizontal_fibra
        total_cobre = self.num_racks_total * 24 * 10 
        
        items.append({"Cat": "Comms", "Item": "Cableado Cobre Cat6A", "Ud": "m", "Cant": total_cobre, "PU": PRECIOS_REF["Cable Cobre Cat6A (m)"]}) 
        items.append({"Cat": "Comms", "Item": "Fibra Óptica (MM/SM)", "Ud": "m", "Cant": total_fibra, "PU": PRECIOS_REF["Fibra Óptica OM4/OS2 (m)"]})
        items.append({"Cat": "Comms", "Item": "Bandejas Fibra/Datos", "Ud": "m", "Cant": dist_lineas_sala, "PU": PRECIOS_REF["Bandeja Rejilla/Fibra (m)"]})
        
        puntos_bms = (n_equipos_hvac * 10) + (lados * 20) + (self.num_racks_total * 2) 
        items.append({"Cat": "BMS", "Item": "Integración BMS/DCIM", "Ud": "Puntos", "Cant": puntos_bms, "PU": PRECIOS_REF["Punto BMS/Integración (ud)"]})
        items.append({"Cat": "Seguridad", "Item": "CCTV & Accesos", "Ud": "Global", "Cant": 1, "PU": (self.cctv_unidades * PRECIOS_REF["Cámara CCTV (ud)"]) + (self.control_accesos_pax * PRECIOS_REF["Control Acceso (punto)"])})

        df = pd.DataFrame(items)
        df["Total (€)"] = df["Cant"] * df["PU"]
        return df

    # --- MÉTODOS DE CÁLCULO PREVIOS ---
    def dimensionar_sistema_hvac_completo(self):
        Q_DLC_capturada_kW = (self.P_IT_demandada * (self.cerramientos_con_dlc / self.num_cerramientos) * self.Eficiencia_Captura_DLC) / 1000
        Q_Total_IT_kW = self.P_IT_demandada / 1000
        Q_Remanente_Aire_kW = Q_Total_IT_kW - Q_DLC_capturada_kW
        factor_ineficiencia = 1.05 if self.tipo_cerramiento == "Pasillo Frío" else 1.25
        Q_HVAC_Diseno_kW = Q_Remanente_Aire_kW * factor_ineficiencia
        Q_Instalada_kW = Q_HVAC_Diseno_kW * self.factor_N_hvac
        capacidad_unitaria = 100.0
        if Q_Instalada_kW > 1000: capacidad_unitaria = 500.0
        
        hidro_prim = self._calcular_tuberia_colector(Q_Instalada_kW, 5.0)
        hidro_sec = self._calcular_tuberia_colector(Q_Instalada_kW, 6.0)

        return {"Q_Diseno_kW": Q_HVAC_Diseno_kW, "Q_Instalada_kW": Q_Instalada_kW, "Hidro_Prim": hidro_prim, "Hidro_Sec": hidro_sec, "Capacidad_Unit": capacidad_unitaria}

    def dimensionar_dlc_hidraulica(self):
        Q_DLC_kW = (self.P_IT_demandada * (self.cerramientos_con_dlc / self.num_cerramientos) * self.Eficiencia_Captura_DLC) / 1000
        hidro_prim = self._calcular_tuberia_colector(Q_DLC_kW, 5.0)
        hidro_sec = self._calcular_tuberia_colector(Q_DLC_kW, 8.0)
        return {"Q_DLC_kW": Q_DLC_kW, "Hidro_Prim": hidro_prim, "Hidro_Sec": hidro_sec}

    def dimensionar_sistema_electrico(self):
        P_Total_N_Watts = self.P_total_demandada 
        S_Total_N_kVA = P_Total_N_Watts / (0.9 * 1000)
        if self.Suministro_AB == "2 Lados (A y B)":
            num_lados = 2; S_Requerida_Por_Lado_kVA = S_Total_N_kVA
        else:
            num_lados = 1; S_Requerida_Por_Lado_kVA = S_Total_N_kVA * self.factor_N_elec

        S_nominal_kVA = [630, 800, 1000, 1250, 1600, 2000, 2500, 3150, 4000]
        T_capacidad = next((s for s in S_nominal_kVA if s >= S_Requerida_Por_Lado_kVA), S_Requerida_Por_Lado_kVA)
        num_celdas_mt = 2 + num_lados
        I_cuadro_IT_A = (T_capacidad * 1000) / (400 * np.sqrt(3))
        P_rack_W = self.servidores_por_rack * self.P_max_servidor
        I_rack_A = (P_rack_W / 400) / np.sqrt(3)
        I_circuito_rack_A = next((i for i in [16, 32, 63, 125] if i >= I_rack_A * 1.25), 32) 
        I_blindobarra_A = next((i for i in [250, 400, 630, 800, 1000, 1250, 1600, 2500, 4000] if i >= I_cuadro_IT_A), I_cuadro_IT_A)
        
        return {"T_capacidad": T_capacidad, "S_Total_N_kVA": S_Total_N_kVA, "I_cuadro_IT": I_cuadro_IT_A, "I_blindobarra": I_blindobarra_A, "I_rack_distribucion": I_circuito_rack_A, "Num_Trafos": num_lados, "Num_Celdas_MT": num_celdas_mt, "Num_Lados": num_lados}

    def calcular_consumos_desglosados(self):
        labels = ['IT', 'HVAC', 'DLC', 'Ilum', 'Control', 'Aux']; sizes = [self.P_IT_demandada, self.P_HVAC_demandada, self.P_DLC_demandada, self.P_iluminacion, self.P_Control_calc, self.P_otras_fuerza + self.P_PCI_calc]
        data = [(labels[i], sizes[i]) for i in range(len(sizes)) if sizes[i] > 1e-3]
        return dict(zip(*zip(*data))) if data else {}

# ==============================================================================
# GENERADORES DE TABLAS (RESTAURADOS EXACTAMENTE)
# ==============================================================================

def generar_tabla_ratios(kpis):
    data = [{"Ratio/KPI": k, "Valor": f"{v:.2f}"} for k, v in kpis.items()]
    return pd.DataFrame(data)

def generar_tabla_electrico(diseno, res):
    T_cap = res['T_capacidad']; Lados = res['Num_Lados']
    data = [
        {"Zona": "Zona 1 (MT)", "Equipo": "Celdas MT Entrada", "Potencia unitaria": "-", "nº de unidades": int(res['Num_Celdas_MT']), "Potencia total": "-", "Nivel de tensión": "24 kV", "Especificaciones": "GIS/AIS SF6"},
        {"Zona": "Zona 2 (Transf)", "Equipo": "Trafo MT/BT", "Potencia unitaria": f"{T_cap} kVA", "nº de unidades": int(Lados), "Potencia total": f"{T_cap * Lados} kVA", "Nivel de tensión": "24kV/400V", "Especificaciones": "Seco/Aceite Dyn11"},
        {"Zona": "Zona 3 (CGBT)", "Equipo": "Cuadro General", "Potencia unitaria": f"{T_cap} kVA", "nº de unidades": int(Lados), "Potencia total": f"{T_cap * Lados} kVA", "Nivel de tensión": "400 V", "Especificaciones": f"In: {res['I_cuadro_IT']:.0f}A, 50kA"},
        {"Zona": "Zona 4 (UPS)", "Equipo": "SAI Modular", "Potencia unitaria": f"{T_cap} kVA", "nº de unidades": int(Lados), "Potencia total": f"{T_cap * Lados} kVA", "Nivel de tensión": "400 V", "Especificaciones": "Doble Conversión"},
        {"Zona": "Zona 5 (Rack)", "Equipo": "Blindobarra/Canalis", "Potencia unitaria": f"{res['I_blindobarra']} A", "nº de unidades": int(diseno.num_cerramientos * Lados), "Potencia total": "-", "Nivel de tensión": "400 V", "Especificaciones": f"Distribución {diseno.Distribucion_IT_tipo}"},
        {"Zona": "Zona 5 (Rack)", "Equipo": "Prot. Circuito Rack", "Potencia unitaria": f"{res['I_rack_distribucion']} A", "nº de unidades": int(diseno.N_servidores_total / diseno.servidores_por_rack * Lados), "Potencia total": f"{diseno.P_IT_demandada/1000:.1f} kW", "Nivel de tensión": "230 V", "Especificaciones": "Magnetotérmico Curva C"}
    ]
    return pd.DataFrame(data)

def generar_tabla_hvac_limpia(diseno, res_hvac):
    Q_tot = res_hvac['Q_Instalada_kW']; Q_unit = res_hvac['Capacidad_Unit']
    N_equipos = np.ceil(Q_tot / Q_unit) if Q_unit > 0 else 1
    data = [
        {"Zona": "Zona 1 (Prod)", "Equipo": diseno.prodfrio_tec, "Potencia unitaria": f"{Q_unit:.0f} kW", "nº de unidades": int(N_equipos), "Potencia total": f"{Q_tot:.0f} kW", "Nivel de tensión": "400V", "Especificaciones": f"Redundancia {diseno.R_hvac}"},
        {"Zona": "Zona 2 (Inter)", "Equipo": diseno.intcalor_tec, "Potencia unitaria": f"{Q_unit:.0f} kW", "nº de unidades": int(N_equipos), "Potencia total": f"{Q_tot:.0f} kW", "Nivel de tensión": "-", "Especificaciones": "Intercambio Térmico"},
        {"Zona": "Zona 3 (Dist)", "Equipo": diseno.distribfrio_tec, "Potencia unitaria": "Var", "nº de unidades": "Var", "Potencia total": f"{Q_tot:.0f} kW", "Nivel de tensión": "230V/400V", "Especificaciones": "Clima Precisión"}
    ]
    return pd.DataFrame(data)

def generar_tabla_hidraulica_unificada(diseno, res_hvac, res_dlc):
    prim_h = res_hvac["Hidro_Prim"]; sec_h = res_hvac["Hidro_Sec"]
    data = [
        {"Zona": "HVAC Primario", "Equipo": "Colector Generación", "Potencia unitaria": "-", "nº de unidades": prim_h['Num_Circuitos'], "Potencia total": "-", "Nivel de tensión": "-", "Especificaciones": f"DN{prim_h['DN_mm']} ({prim_h['Material']}). Q={prim_h['Caudal_Total_m3h']:.1f} m3/h"},
        {"Zona": "HVAC Secundario", "Equipo": "Anillo Sala", "Potencia unitaria": "-", "nº de unidades": sec_h['Num_Circuitos'], "Potencia total": "-", "Nivel de tensión": "-", "Especificaciones": f"DN{sec_h['DN_mm']} ({sec_h['Material']}). V={sec_h['Velocidad_ms']:.2f} m/s"}
    ]
    if diseno.cerramientos_con_dlc > 0:
        prim_d = res_dlc["Hidro_Prim"]; sec_d = res_dlc["Hidro_Sec"]
        data.append({"Zona": "DLC Primario", "Equipo": "Loop Enfriamiento", "Potencia unitaria": "-", "nº de unidades": prim_d['Num_Circuitos'], "Potencia total": "-", "Nivel de tensión": "-", "Especificaciones": f"DN{prim_d['DN_mm']} ({prim_d['Material']}). Q={prim_d['Caudal_Total_m3h']:.1f} m3/h"})
        data.append({"Zona": "DLC Secundario", "Equipo": "Loop Chips", "Potencia unitaria": "-", "nº de unidades": sec_d['Num_Circuitos'], "Potencia total": "-", "Nivel de tensión": "-", "Especificaciones": f"DN{sec_d['DN_mm']} ({sec_d['Material']}). V={sec_d['Velocidad_ms']:.2f} m/s"})
    return pd.DataFrame(data)

def generar_tabla_pci(diseno):
    return pd.DataFrame([
        {"Zona": "General", "Equipo": "Detección + Extinción", "Potencia unitaria": "-", "nº de unidades": int(diseno.centralitas_incendios), "Potencia total": f"{diseno.P_PCI_calc} W", "Nivel de tensión": "230V", "Especificaciones": f"Tecnología: {diseno.tecnologia_pci}"},
        {"Zona": "Sala Bombas", "Equipo": "Grupo Presión", "Potencia unitaria": "20kW", "nº de unidades": int(diseno.grupos_bombeo_pci), "Potencia total": f"{diseno.grupos_bombeo_pci*20} kW", "Nivel de tensión": "400V", "Especificaciones": "Bomba Ppal + Reserva"}
    ])

def generar_tabla_control(diseno):
    return pd.DataFrame([
        {"Zona": "Seguridad", "Equipo": "CCTV & Accesos", "Potencia unitaria": "-", "nº de unidades": int(diseno.cctv_unidades), "Potencia total": f"{diseno.P_Control_calc} W", "Nivel de tensión": "PoE", "Especificaciones": f"Cámaras: {diseno.cctv_unidades}, Puntos Acc: {diseno.control_accesos_pax}"}
    ])

# ==============================================================================
# GRÁFICOS (RESTAURADOS)
# ==============================================================================
def generar_grafico_metricas(diseno, WCR, CEF):
    if diseno.P_IT_demandada > 0:
        PUE = diseno.P_total_demandada / diseno.P_IT_demandada
        WUE = (diseno.P_HVAC_demandada / diseno.P_IT_demandada) * WCR 
        CUE = PUE * CEF
    else:
        PUE = 1.0; WUE = 0.0; CUE = 0.0

    metrics = [PUE, CUE, WUE]
    names = ['PUE (Ratio)', f'CUE (kgCO2/kWh)', f'WUE (L/kWh)']
    colors = ['#FF6F61', '#6B5B95', '#88B04B']

    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.bar(names, metrics, color=colors)
    ax.axhline(y=1.0, color='gray', linestyle='--', alpha=0.7)
    ax.set_title('KPIs de Eficiencia y Sostenibilidad')
    ax.set_ylabel('Valor')
    
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'{height:.2f}', xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
    
    return fig

def generar_grafico_consumos(consumos):
    if not consumos: return None
    labels = list(consumos.keys())
    sizes = list(consumos.values())
    colors = ['#4CAF50', '#2196F3', '#FFC107', '#9E9E9E', '#607D8B', '#FF5722']
    
    fig, ax = plt.subplots(figsize=(6, 6))
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%', 
                                      startangle=90, colors=colors[:len(sizes)], 
                                      pctdistance=0.85, wedgeprops=dict(width=0.4))
    plt.setp(texts, size=10)
    plt.setp(autotexts, size=10, weight="bold", color="white")
    ax.set_title("Desglose de Consumo Energético")
    centre_circle = plt.Circle((0,0),0.70,fc='white')
    fig.gca().add_artist(centre_circle)
    ax.axis('equal')  
    return fig

# ==============================================================================
# GENERACIÓN DE REPORTE WORD (RESTAURADA EXACTA)
# ==============================================================================
def crear_documento_proyecto_word(diseno, df_elec, df_hvac, df_hidro, df_pci, consumos, df_capex, df_ratios, fig_consumos, fig_metricas):
    if not HAS_DOCX: return None
    
    doc = Document()
    
    # --- PORTADA ---
    doc.add_heading('PROYECTO EJECUTIVO DE DATA CENTER', 0)
    doc.add_paragraph(f'Fecha de Generación: {datetime.date.today().strftime("%d/%m/%Y")}')
    doc.add_paragraph('Este documento contiene la memoria técnica descriptiva, los cálculos justificativos y el presupuesto estimado para la infraestructura del CPD.')
    doc.add_page_break()
    
    # --- 1. INTRODUCCIÓN ---
    doc.add_heading('1. Introducción y Objeto', level=1)
    p = doc.add_paragraph()
    p.add_run(f'El objeto del presente proyecto es definir las instalaciones de un Centro de Procesamiento de Datos (CPD) con una superficie total construida de {diseno.area_total_construida:.0f} m² distribuidos en {diseno.num_plantas} plantas. ')
    p.add_run(f'La infraestructura dará servicio a una carga IT crítica de {diseno.P_IT_demandada/1000:.2f} kW, alojada en {diseno.num_racks_total} racks de servidores.')
    
    # --- 2. RATIOS Y MÉTRICAS ---
    doc.add_heading('2. Ratios de Diseño y Eficiencia', level=1)
    doc.add_paragraph('A continuación se detallan los indicadores clave de rendimiento (KPIs) calculados para el diseño propuesto:')
    
    if fig_metricas:
        memfile = BytesIO()
        fig_metricas.savefig(memfile, format='png', bbox_inches='tight')
        doc.add_picture(memfile, width=Inches(5))
        memfile.close()
    
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Indicador"
    t.rows[0].cells[1].text = "Valor"
    for _, row in df_ratios.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(row["Ratio/KPI"])
        cells[1].text = str(row["Valor"])
        
    doc.add_page_break()

    # --- 3. INSTALACIÓN ELÉCTRICA ---
    doc.add_heading('3. Instalación Eléctrica', level=1)
    
    desc_elec = (f"El sistema eléctrico se ha diseñado bajo una topología de redundancia {diseno.R_elec}, con un esquema de suministro "
                 f"tipo {diseno.Suministro_AB}. La distribución de potencia hacia la sala IT se realizará mediante {diseno.Distribucion_IT_tipo}. "
                 f"Se contempla la instalación de {df_elec.iloc[2]['nº de unidades']} transformadores de {df_elec.iloc[2]['Potencia unitaria']} "
                 f"y un sistema de respaldo mediante grupos electrógenos y sistemas de alimentación ininterrumpida (UPS) de doble conversión.")
    doc.add_paragraph(desc_elec)
    
    doc.add_heading('3.1. Equipos Eléctricos Principales', level=2)
    t = doc.add_table(rows=1, cols=len(df_elec.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df_elec.columns): t.rows[0].cells[i].text = col
    for _, row in df_elec.iterrows():
        rc = t.add_row().cells
        for i, val in enumerate(row): rc[i].text = str(val)

    # --- 4. CLIMATIZACIÓN ---
    doc.add_heading('4. Sistema HVAC', level=1)
    
    desc_hvac = (f"La disipación térmica se gestionará mediante un sistema de producción de frío basado en tecnología {diseno.prodfrio_tec} "
                 f"con redundancia {diseno.R_hvac}. La distribución de aire en sala se realizará mediante unidades {diseno.distribfrio_tec} "
                 f"configuradas para un cerramiento de {diseno.tipo_cerramiento}. ")
    
    if diseno.cerramientos_con_dlc > 0:
        desc_hvac += (f"Adicionalmente, se implementa un sistema de Refrigeración Líquida Directa (DLC) de alta densidad para {diseno.cerramientos_con_dlc} cerramientos, "
                      f"utilizando {diseno.tipo_dist_frio_dlc} y generación mediante {diseno.tipo_gen_frio_dlc}.")
    
    doc.add_paragraph(desc_hvac)
    
    doc.add_heading('4.1. Equipos de Climatización', level=2)
    t = doc.add_table(rows=1, cols=len(df_hvac.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df_hvac.columns): t.rows[0].cells[i].text = col
    for _, row in df_hvac.iterrows():
        rc = t.add_row().cells
        for i, val in enumerate(row): rc[i].text = str(val)

    # --- 5. HIDRÁULICA ---
    doc.add_heading('5. Red Hidráulica', level=1)
    doc.add_paragraph("A continuación se detallan las características de los circuitos hidráulicos calculados (diámetros, materiales y caudales) para garantizar el transporte de energía térmica:")
    
    t = doc.add_table(rows=1, cols=len(df_hidro.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df_hidro.columns): t.rows[0].cells[i].text = col
    for _, row in df_hidro.iterrows():
        rc = t.add_row().cells
        for i, val in enumerate(row): rc[i].text = str(val)

    # --- 6. PCI Y SEGURIDAD ---
    doc.add_heading('6. Protección Contra Incendios y Seguridad', level=1)
    doc.add_paragraph(f"El sistema de extinción seleccionado para las salas críticas es {diseno.tecnologia_pci}, complementado por un sistema de detección temprana VESDA. La seguridad física se gestiona mediante un sistema integrado de CCTV y control de accesos.")
    
    t = doc.add_table(rows=1, cols=len(df_pci.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df_pci.columns): t.rows[0].cells[i].text = col
    for _, row in df_pci.iterrows():
        rc = t.add_row().cells
        for i, val in enumerate(row): rc[i].text = str(val)

    doc.add_page_break()

    # --- 7. ANÁLISIS DE CONSUMOS ---
    doc.add_heading('7. Análisis Energético', level=1)
    doc.add_paragraph("Desglose estimado de la potencia demandada por subsistema:")
    
    if fig_consumos:
        memfile = BytesIO()
        fig_consumos.savefig(memfile, format='png', bbox_inches='tight')
        doc.add_picture(memfile, width=Inches(5))
        memfile.close()
        
    doc.add_paragraph("Detalle de Potencias (W):")
    for k, v in consumos.items():
        doc.add_paragraph(f"- {k}: {v:.0f} W", style='List Bullet')

    # --- 8. PRESUPUESTO ---
    doc.add_heading('8. Presupuesto Estimado (CAPEX)', level=1)
    doc.add_paragraph("Estimación de costes de ejecución material (PEM) basada en precios de mercado de referencia:")
    
    t = doc.add_table(rows=1, cols=len(df_capex.columns))
    t.style = 'Table Grid'
    for i, col in enumerate(df_capex.columns): t.rows[0].cells[i].text = col
    for _, row in df_capex.iterrows():
        rc = t.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float):
                rc[i].text = f"{val:,.2f}"
            else:
                rc[i].text = str(val)
    
    total_capex = df_capex['Total (€)'].sum()
    doc.add_paragraph(f"\nTOTAL ESTIMADO: {total_capex:,.2f} €", style='Heading 2')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================================================================
# GUI DE ESCRITORIO (TKINTER) - CONECTANDO TODO
# ==============================================================================
class DesktopCPDApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Ingeniería CPD v15.1 - Desktop Edition")
        self.root.geometry("1400x900")
        
        # --- Variables de Entrada (Inputs) ---
        self.vars = {
            "num_plantas": tk.IntVar(value=2),
            "area_planta": tk.DoubleVar(value=500.0),
            "area_it": tk.DoubleVar(value=400.0),
            "num_cerramientos": tk.IntVar(value=4),
            "racks_por_cerramiento": tk.IntVar(value=12),
            "servidores_por_rack": tk.IntVar(value=10),
            "P_max": tk.DoubleVar(value=500.0),
            "P_idle": tk.DoubleVar(value=100.0),
            "red_elec": tk.StringVar(value="2N"),
            "suministro_AB": tk.StringVar(value="2 Lados (A y B)"),
            "dist_it": tk.StringVar(value="Blindobarra"),
            "cop_hvac": tk.DoubleVar(value=3.5),
            "t_in": tk.DoubleVar(value=22.0),
            "t_out": tk.DoubleVar(value=34.0),
            "p_ilum": tk.DoubleVar(value=2000.0),
            "tipo_cerr": tk.StringVar(value="Pasillo Frío"),
            "prod_frio": tk.StringVar(value="Chiller A/W"),
            "int_calor": tk.StringVar(value="Placas Soldadas"),
            "dist_frio": tk.StringVar(value="CRAH"),
            "tec_pci": tk.StringVar(value="Agua Nebulizada"),
            "cent_pci": tk.IntVar(value=2),
            "vesda": tk.IntVar(value=4),
            "bombas": tk.IntVar(value=1),
            "cctv": tk.IntVar(value=20),
            "accesos": tk.IntVar(value=10),
            "WCR": tk.DoubleVar(value=0.5),
            "CEF": tk.DoubleVar(value=0.35),
            "n_dlc": tk.IntVar(value=0),
            "eff_dlc": tk.DoubleVar(value=0.8),
            "gen_dlc": tk.StringVar(value="Dry cooler adiabático"),
            "dist_dlc": tk.StringVar(value="CDU in-rack"),
            "cop_dlc": tk.DoubleVar(value=10.0),
            "aux_dlc": tk.DoubleVar(value=500.0)
        }

        # --- Layout Principal ---
        main_frame = ttk.Frame(root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Panel izquierdo: Inputs
        left_panel = ttk.Frame(main_frame, width=400)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        
        # Notebook de Inputs
        input_tabs = ttk.Notebook(left_panel)
        input_tabs.pack(fill=tk.BOTH, expand=True)
        
        self.create_geo_tab(input_tabs)
        self.create_clima_tab(input_tabs)
        self.create_equip_tab(input_tabs)
        self.create_dlc_tab(input_tabs)

        # Botón Calcular
        calc_btn = ttk.Button(left_panel, text="▶ CALCULAR PROYECTO", command=self.run_calculation)
        calc_btn.pack(fill=tk.X, pady=10)
        
        # Botón Exportar
        self.export_btn = ttk.Button(left_panel, text="📥 EXPORTAR DOCX", command=self.export_report, state=tk.DISABLED)
        self.export_btn.pack(fill=tk.X)

        # Panel derecho: Resultados
        self.right_panel = ttk.Notebook(main_frame)
        self.right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # Pestañas de resultados
        self.tab_kpi = ttk.Frame(self.right_panel); self.right_panel.add(self.tab_kpi, text="KPIs & Gráficos")
        self.tab_capex = ttk.Frame(self.right_panel); self.right_panel.add(self.tab_capex, text="Presupuesto (CAPEX)")
        self.tab_elec = ttk.Frame(self.right_panel); self.right_panel.add(self.tab_elec, text="Electricidad")
        self.tab_hvac = ttk.Frame(self.right_panel); self.right_panel.add(self.tab_hvac, text="Mecánica")
        self.tab_aux = ttk.Frame(self.right_panel); self.right_panel.add(self.tab_aux, text="Auxiliares")

        # Variables para almacenar resultados
        self.current_design = None
        self.current_dfs = {}
        self.current_figs = {}
        self.current_consumos = {}

    # --- Helpers para Inputs ---
    def add_entry(self, parent, label, var, r):
        ttk.Label(parent, text=label).grid(row=r, column=0, sticky="w", pady=2)
        ttk.Entry(parent, textvariable=var, width=15).grid(row=r, column=1, sticky="e", pady=2)

    def add_combo(self, parent, label, var, values, r):
        ttk.Label(parent, text=label).grid(row=r, column=0, sticky="w", pady=2)
        ttk.Combobox(parent, textvariable=var, values=values, width=13, state="readonly").grid(row=r, column=1, sticky="e", pady=2)

    def create_geo_tab(self, notebook):
        frame = ttk.Frame(notebook, padding=10)
        notebook.add(frame, text="Geometría")
        self.add_entry(frame, "Nº Plantas:", self.vars["num_plantas"], 0)
        self.add_entry(frame, "Area Planta (m²):", self.vars["area_planta"], 1)
        self.add_entry(frame, "Area Sala IT (m²):", self.vars["area_it"], 2)
        ttk.Separator(frame, orient=tk.HORIZONTAL).grid(row=3, columnspan=2, sticky="ew", pady=5)
        self.add_entry(frame, "Nº Cerramientos:", self.vars["num_cerramientos"], 4)
        self.add_entry(frame, "Racks/Cerramiento:", self.vars["racks_por_cerramiento"], 5)
        self.add_entry(frame, "Servers/Rack:", self.vars["servidores_por_rack"], 6)
        self.add_entry(frame, "W/Server (Max):", self.vars["P_max"], 7)
        self.add_combo(frame, "Redundancia Elec:", self.vars["red_elec"], ["2N", "N+1", "N"], 8)

    def create_clima_tab(self, notebook):
        frame = ttk.Frame(notebook, padding=10)
        notebook.add(frame, text="Clima/Elec")
        self.add_combo(frame, "Suministro:", self.vars["suministro_AB"], ["2 Lados (A y B)", "1 Lado (A)"], 0)
        self.add_combo(frame, "Distrib. BT:", self.vars["dist_it"], ["Blindobarra", "Cable"], 1)
        self.add_entry(frame, "COP HVAC:", self.vars["cop_hvac"], 2)
        self.add_entry(frame, "T Entrada (°C):", self.vars["t_in"], 3)
        self.add_entry(frame, "T Salida (°C):", self.vars["t_out"], 4)
        self.add_entry(frame, "Iluminación (W):", self.vars["p_ilum"], 5)

    def create_equip_tab(self, notebook):
        frame = ttk.Frame(notebook, padding=10)
        notebook.add(frame, text="Equipos")
        self.add_combo(frame, "Tipo Cerramiento:", self.vars["tipo_cerr"], ["Pasillo Frío", "Pasillo Caliente", "Sin Cerramiento"], 0)
        self.add_combo(frame, "Prod. Frío:", self.vars["prod_frio"], ["Condensadora DX", "Chiller A/W", "Chiller A/W con free cooling", "Chiller W/W", "Dry cooler seco", "Chiller W/W + Torre de refrigeración", "Torre de refrigeración"], 1)
        self.add_combo(frame, "Intercambio:", self.vars["int_calor"], ["Placas Soldadas", "Tubular", "Ninguno (Directo)"], 2)
        self.add_combo(frame, "Distrib. Frío:", self.vars["dist_frio"], ["CRAH", "CRAC", "Inrow agua", "Inrow DX", "Puerta trasera RDHx", "Inmersión en dieléctrico", "CDU central", "CDU in-row", "CDU in-rack"], 3)
        ttk.Separator(frame, orient=tk.HORIZONTAL).grid(row=4, columnspan=2, sticky="ew", pady=5)
        self.add_combo(frame, "Extinción PCI:", self.vars["tec_pci"], ["Agua Nebulizada", "NOVEC 1230", "ARGONITE", "FM-200"], 5)
        self.add_entry(frame, "Centralitas PCI:", self.vars["cent_pci"], 6)
        self.add_entry(frame, "VESDA:", self.vars["vesda"], 7)
        self.add_entry(frame, "Bombas PCI:", self.vars["bombas"], 8)
        self.add_entry(frame, "Cámaras CCTV:", self.vars["cctv"], 9)
        self.add_entry(frame, "Accesos:", self.vars["accesos"], 10)

    def create_dlc_tab(self, notebook):
        frame = ttk.Frame(notebook, padding=10)
        notebook.add(frame, text="DLC/Sustain")
        self.add_entry(frame, "Cerramientos DLC:", self.vars["n_dlc"], 0)
        self.add_entry(frame, "Efic. Captura (0-1):", self.vars["eff_dlc"], 1)
        self.add_combo(frame, "Generación DLC:", self.vars["gen_dlc"], ["Dry cooler adiabático", "Chiller A/W de alta temperatura", "Torre de refrigeración"], 2)
        self.add_combo(frame, "Distribución DLC:", self.vars["dist_dlc"], ["CDU central", "CDU in-row", "CDU in-rack", "Inmersión en dieléctrico"], 3)
        self.add_entry(frame, "COP DLC:", self.vars["cop_dlc"], 4)
        self.add_entry(frame, "Pot Aux DLC (W):", self.vars["aux_dlc"], 5)
        ttk.Separator(frame, orient=tk.HORIZONTAL).grid(row=6, columnspan=2, sticky="ew", pady=5)
        self.add_entry(frame, "WCR:", self.vars["WCR"], 7)
        self.add_entry(frame, "CEF:", self.vars["CEF"], 8)

    # --- Lógica de Ejecución ---
    def run_calculation(self):
        try:
            # 1. Instanciar Motor
            self.current_design = DisenadorV14(
                self.vars["red_elec"].get(), "N+1", self.vars["suministro_AB"].get(), self.vars["dist_it"].get(),
                self.vars["num_cerramientos"].get(), self.vars["racks_por_cerramiento"].get(), self.vars["servidores_por_rack"].get(),
                self.vars["tipo_cerr"].get(), self.vars["P_idle"].get(), self.vars["P_max"].get(),
                self.vars["p_ilum"].get(), 3000, self.vars["cop_hvac"].get(), self.vars["t_in"].get(), self.vars["t_out"].get(),
                self.vars["prod_frio"].get(), self.vars["int_calor"].get(), self.vars["dist_frio"].get(), 2,
                self.vars["n_dlc"].get(), self.vars["gen_dlc"].get(), self.vars["cop_dlc"].get(),
                self.vars["dist_dlc"].get(), self.vars["aux_dlc"].get(), self.vars["eff_dlc"].get(),
                self.vars["cent_pci"].get(), self.vars["vesda"].get(), self.vars["bombas"].get(),
                self.vars["cctv"].get(), self.vars["accesos"].get(), self.vars["tec_pci"].get(),
                self.vars["num_plantas"].get(), self.vars["area_planta"].get(), self.vars["area_it"].get()
            )

            # 2. Correr Cálculos
            res_elec = self.current_design.dimensionar_sistema_electrico()
            res_hvac = self.current_design.dimensionar_sistema_hvac_completo()
            res_dlc = self.current_design.dimensionar_dlc_hidraulica()
            self.current_consumos = self.current_design.calcular_consumos_desglosados()
            
            # 3. Generar DataFrames
            df_capex = self.current_design.calcular_presupuesto_detallado(res_elec, res_hvac, res_dlc)
            kpis = self.current_design.calcular_kpis_densidad(res_hvac['Q_Instalada_kW'], res_elec['S_Total_N_kVA'])
            
            # Tablas para GUI (Usando los generadores restaurados)
            df_elec_t = generar_tabla_electrico(self.current_design, res_elec)
            df_hvac_t = generar_tabla_hvac_limpia(self.current_design, res_hvac)
            df_hidro_t = generar_tabla_hidraulica_unificada(self.current_design, res_hvac, res_dlc)
            df_pci_t = pd.concat([generar_tabla_pci(self.current_design), generar_tabla_control(self.current_design)])
            df_ratios_t = generar_tabla_ratios(kpis)

            # Guardar para exportación
            self.current_dfs = {
                "capex": df_capex, "elec": df_elec_t, "hvac": df_hvac_t, 
                "hidro": df_hidro_t, "pci": df_pci_t, "ratios": df_ratios_t
            }

            # 4. Actualizar GUI
            self.render_dataframe(self.tab_capex, df_capex)
            self.render_dataframe(self.tab_elec, df_elec_t)
            self.render_dataframe(self.tab_hvac, pd.concat([df_hvac_t, df_hidro_t]))
            self.render_dataframe(self.tab_aux, df_pci_t)
            
            # Renderizar KPIs (Gráficos + Tabla)
            self.render_kpi_tab(kpis, self.current_consumos)

            self.export_btn.config(state=tk.NORMAL)
            messagebox.showinfo("Cálculo Exitoso", f"Inversión Estimada: {df_capex['Total (€)'].sum():,.2f} €")

        except Exception as e:
            messagebox.showerror("Error en Cálculo", str(e))

    def render_dataframe(self, parent_widget, df):
        # Limpiar
        for widget in parent_widget.winfo_children(): widget.destroy()
        
        # Treeview
        columns = list(df.columns)
        tree = ttk.Treeview(parent_widget, columns=columns, show="headings")
        
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=120)  # Ajustar ancho
        
        for index, row in df.iterrows():
            vals = []
            for v in row:
                if isinstance(v, float): vals.append(f"{v:,.2f}")
                else: vals.append(v)
            tree.insert("", tk.END, values=vals)
        
        # Scrollbars
        vsb = ttk.Scrollbar(parent_widget, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(parent_widget, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb.pack(side=tk.BOTTOM, fill=tk.X)

    def render_kpi_tab(self, kpis, consumos):
        for widget in self.tab_kpi.winfo_children(): widget.destroy()
        
        # Frame superior para gráficos
        graph_frame = ttk.Frame(self.tab_kpi)
        graph_frame.pack(fill=tk.BOTH, expand=True)
        
        # Gráfico Métricas
        fig1 = generar_grafico_metricas(self.current_design, self.vars["WCR"].get(), self.vars["CEF"].get())
        canvas1 = FigureCanvasTkAgg(fig1, master=graph_frame)
        canvas1.draw()
        canvas1.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.current_figs["metricas"] = fig1

        # Gráfico Consumos
        fig2 = generar_grafico_consumos(consumos)
        if fig2:
            canvas2 = FigureCanvasTkAgg(fig2, master=graph_frame)
            canvas2.draw()
            canvas2.get_tk_widget().pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
            self.current_figs["consumos"] = fig2

        # Frame inferior para tabla
        table_frame = ttk.Frame(self.tab_kpi, height=150)
        table_frame.pack(fill=tk.X)
        self.render_dataframe(table_frame, generar_tabla_ratios(kpis))

    def export_report(self):
        if not HAS_DOCX:
            messagebox.showwarning("Falta Librería", "Instala 'python-docx' para exportar.")
            return

        filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if filename:
            try:
                # LLAMADA A LA FUNCIÓN ORIGINAL RESTAURADA
                doc_buffer = crear_documento_proyecto_word(
                    self.current_design, 
                    self.current_dfs["elec"], 
                    self.current_dfs["hvac"],
                    self.current_dfs["hidro"], 
                    self.current_dfs["pci"], 
                    self.current_consumos,
                    self.current_dfs["capex"], 
                    self.current_dfs["ratios"], 
                    self.current_figs.get("consumos"), 
                    self.current_figs.get("metricas")
                )
                with open(filename, "wb") as f:
                    f.write(doc_buffer.getbuffer())
                messagebox.showinfo("Exportar", "Informe generado correctamente.")
            except Exception as e:
                messagebox.showerror("Error Exportando", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    style = ttk.Style()
    style.theme_use('clam') 
    app = DesktopCPDApp(root)
    root.mainloop()
